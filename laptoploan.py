#  user inputs student name, device SN / asset #, and device type
# this py file will use two word docs as originals and create a new laptop loan 
# agreement file and welcome letter with the corresponding details included. Possibly 
# send to device's default printer afterwards
import docx # pip install python-docx
from docx.shared import Pt 
import os
import traceback
import pyperclip
import re # for error checking

# ** DISCLAIMER ** --> This script WILL attempt to print the documents to your default printer if it runs w/o errors

# font for inserted text
FONT = "Times New Roman"
# Location inserted into the text copied to clipboard (more useful during OsTicket email issue)
LOCATION = "Stanton campus Open Computer Lab - room A227"

class LaptopLoan:
    def __init__(self, thefont, location):
        self.font = thefont
        self.location = location
        # path to the welcome / agreement base files
        self.welcome_file = './files/welcome.docx'
        self.agreement = './files/agreement.docx'

        #  create the Document objects from the two word docs (I converted the fillable PDF to a word doc bc they're easier to edit w/python)
        try:
            self.welcome_doc = docx.Document(self.welcome_file)
            self.agree_doc = docx.Document(self.agreement)
        except Exception:
            # if it fails here - probably because it didn't find the welcome/agree files
            # print(traceback.format_exc())

            print(f"Please make sure you have the {self.welcome_file} and {self.agreement} files stored in {os.path.basename(os.getcwd())}")
            # no point in continuing script w/o the files
            exit()

        # start loop to allow for multiple laptop loan application processing
        while True:
            # ask user for student's name:
            name = input('Student\'s name: ')
            # service tag (ST) - 'serial number' printed to console bc if it was 'service tag' it could easily be confused with 'asset tag' 
            sn = input('Serial number (ST): ')
            asset = input('Asset Tag: ')
            # one method for error checking the asset tag entry to make sure it has 7 digits and starts with 100-900
            tries = 0
            while True:
                if re.match(r"\d\d\d\d\d\d\d", asset):
                    break # break out of loop if asset tag entry matches the above regex pattern
                # if the user has tried > 7 times to enter asset, add frown emoticon and directional arrow to request
                elif tries < 7:
                    asset = input('Asset Tag: ')
                    tries += 1
                else:
                    asset = input('Asset Tag :/ --> ')
                    tries += 1

            device = input('Make/model of device: ')

            # -- Create the Welcome Letter document
            try:
                self.create_welcome(name)
            except Exception:
                #print(traceback.format_exc())
                print('\nOh no, something went wrong with creating the welcome document!')
                exit()
            # -- Create the Laptop Loan Agreement document
            try:
                self.create_agreement(name, sn, asset, device)
                # print("agreement created")
            except Exception:
                #print(traceback.format_exc())
                print('\nOh no, something went wrong with creating the agreement document!')
                exit()

            # print msg to console saying that files have been created
            #print(f'Sending {name}\'s documents to default printer...')
            try:
                # print files
                self.print_em()
            # except for any error, then..
            except Exception:
                # print(traceback.format_exc())
                print(f"Something went wrong with printing the documents, but they should be saved under the student's name in the {os.path.basename(os.getcwd())} directory.")
                exit()
            # generate the email to send to student letting them know about laptop
            print('Generating text for notification email to student...')

            # email text will be generated with the student's name, IT person can change the 'location' variable to wherever they want student to go (i.e. at Stanton we'd say Stanton Open Lab A227 or something)
            self.gen_email_text(name)

            go_again = input('Print docs for another user? (y/n): ')
            while go_again.lower() not in ['y', 'n']:
                go_again = input('Please enter Y to enter another user\'s info, or N to kill program')

            # the only way program reaches this point is if the go_again variable equals y or n
            # if its y - repeat while loop
            if go_again.lower() == "y":
                pass
            # if go_again = anything except "y" (like "n"), exit program
            else:
                break

    # user inputs name, then a word doc is created with students name inserted
    def create_welcome(self, student_name):
        # grab first sentence of the body/text of the base welcome doc
        first_sentence = self.welcome_doc.paragraphs[1]
        # The only variable text on the document
        welcome_string = f"Dear {student_name[:-12]}"
        # rest of that chunk of text on the document
        rest_of_welcome = ": Welcome to Delaware Technical Community College!  This letter contains some helpful information pertaining to your laptop’s configuration of software and virtual support for assisting with technology issues."

        # replace the <<student-name>> string with the student's name
        first_sentence.text = "" # erase the first sentence
        # item, item_font, item_size, bold, underline
        greeting = first_sentence.add_run(welcome_string) # add the new welcome string into the blank first sentence, and save the new 
        self.format_item(greeting,Pt(12), True, False) # sentence/python docx object to greeting variable so it can be used
                                                       # in the formatting function
        # then add the rest of the paragraph, not bold
        rest = first_sentence.add_run(rest_of_welcome)
        self.format_item(rest, Pt(12), False, False)

        # save the welcome file
        self.welcome_filepath = f'{student_name}-welcome.docx'
        self.welcome_doc.save(self.welcome_filepath)

    # the doc creation scripts take out the chunk of the document that needs to be changed, then inserts the chunk back in w/the variable parts input by user
    def create_agreement(self, student_name, servtag, assettag, dev):

        # deal with the first sentence - insert student's name
        first_sentence = self.agree_doc.paragraphs[2]
        start = "On the date of _____/_____/______, I, "

        first_sentence.text = ""
        started = first_sentence.add_run(start)
        self.format_item(started, Pt(12), False, False)
        # making the students name a bit bigger on the form, the different formatting is why there are 3 add_runs here instead of just one to add the whole first sentence (could be quicker way to do this)
        add_name = first_sentence.add_run(f'{student_name}')
        self.format_item(add_name, Pt(13), False, True)
        add_rest = first_sentence.add_run(', received the following computer equipment and accessories (“the Equipment”) from Delaware Technical Community College (“DTCC”):')
        self.format_item(add_rest, Pt(12), False, False)

        # inserting the device make/model, service tag and asset #
        two_underlines = self.agree_doc.paragraphs[3] # first have to target the chunk of text and erase it
        two_underlines.text = ""
        # add some blank space before the underlined make/model part starts
        two_underlines.add_run('         ')
        # check how many characters are in make/model string
        devlength = len(dev)
        # 56 spaces is the total number of spaces - underlines, make/model string included from start to end
        # subtracting the length of dev string input by user from 56, then dividing by 2 and rounding to whole number gives you the amount of underline you should have on either side of the device string to make it look even
        spaces = " " * round(((56 - devlength)/2))
        # add to document, spaces included
        make_model = two_underlines.add_run(f'{spaces}{dev}{spaces}')
        self.format_item(make_model, Pt(13), False, True)
        # put a tab of space in between the device description and the serial/asset
        two_underlines.add_run('\t')
        # format: <serial number> / <asset tag>, ex: J854MNB / 2003198
        assetstring = f"{servtag} / {assettag}"
        tags = two_underlines.add_run(f'{assetstring}')
        self.format_item(tags, Pt(13),False, True)

        # create new agreement file with student's name
        self.agreement_filepath = f'{student_name}-agreement.docx'
        self.agree_doc.save(self.agreement_filepath)

    def print_em(self):
        # print the welcome letter to default printer (only Windows compatible according to stackoverflow):
        os.startfile(self.welcome_filepath, "print")
        # print agreement
        os.startfile(self.agreement_filepath, "print")
    
    # this function was more useful when OSTicket email responses were being maintenanced
    def gen_email_text(self, student_name):
        # copy email text to clipboard automatically as well, so if people use clipboard history then they can copy multiple/different student's emails to clipboard in one instance of the program:
        print("Copying email text to clipboard...")
        pyperclip.copy(f'Hi {student_name},\nYour laptop is available for pickup at your earliest convenience (within next 5 business days) at: {self.location}. If you have any questions, you can respond to this email or call the college-wide IT Help Desk line: 302-857-1700. Hope to see you soon!')

    # function to format text so not just typing these formatting lines over and over for each piece of text
    def format_item(self, item, item_size, bold, underline):
        item.font.name = self.font
        item.font.size = item_size
        item.bold = bold    # bold isnt showing on document when done this way
        item.underline = underline


if __name__ == '__main__':
    loan = LaptopLoan(FONT, LOCATION)